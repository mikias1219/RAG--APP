"""Extract plain text from supported file types for RAG indexing."""

from __future__ import annotations

from pathlib import Path


def extract_text_from_path(path: Path | str, original_name: str | None = None) -> str:
    """
    Return UTF-8 text from .txt, .md, .pdf, or .docx.
    Raises ValueError with a short message on unsupported type or empty extraction.
    """
    p = Path(path)
    name = (original_name or p.name).lower()
    suffix = Path(name).suffix.lower()

    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="replace").strip()

    if suffix == ".pdf":
        return _text_from_pdf(p)

    if suffix == ".docx":
        return _text_from_docx(p)

    raise ValueError(f"Unsupported file type: {suffix or '(none)'}")


def _text_from_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Install pypdf: pip install pypdf") from exc

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError("No text could be extracted from this PDF (image-only or protected).")
    return text


def _text_from_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise RuntimeError("Install python-docx: pip install python-docx") from exc

    document = docx.Document(str(path))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError("No text could be extracted from this Word document.")
    return text


def allowed_upload_suffixes() -> frozenset[str]:
    return frozenset({".txt", ".md", ".pdf", ".docx"})
